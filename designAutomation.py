## there's a lot of work to be done here

from _cns import *
from _st import *
from PIL import Image
from PIL import ImageOps
import sys,os,shutil,re,json,pysftp,paramiko
from docx.api import Document

def md(dn):
    os.mkdir(dn)
    os.mkdir(lv)
    os.mkdir(p+'ar\\'+dn)

# def mf(dn,**GTag=False,**CTag=False):
#     dn=__import__(dn+'_d',fromlist=['i','l'])
#     locals().update(i|l)
#     if pgs==[]:
#         with open(lb+'m.html') as fi,open(dd+'index.html','w') as f:
#             for l in fi.readlines():
#                 ls=l.lstrip()
#                 if ls.startswith('<meta name="viewport"'):
#                     f.write(l)
#                     f.write(t2+'<title>'+bn+'</title>\n')
#                     f.write(t2+'<base href="'+base+'" target="_self">\n')
#                     continue
#                 elif ls.startswith('<link'):
#                     f.write(l)
#                     if GTag!=False:
#                         f.write(t2+'<script async src="https://www.googletagmanager.com/gtag/js?id=G-'+GTag+'"></script><script>window.dataLayer=window.dataLayer || []; function gtag(){dataLayer.push(arguments);} gtag(\'js\',new Date()); gtag(\'config\',\'G-'+GTag+'\');</script>\n')
#                     if CTag!=False:
#                         f.write(t2+'<script>(function(c,l,a,r,i,t,y){c[a]=c[a]||function(){(c[a].q=c[a].q||[]).push(arguments)};t=l.createElement(r);t.async=1;t.src="https://www.clarity.ms/tag/"+i;y=l.getElementsByTagName(r)[0];y.parentNode.insertBefore(t,y);})(window,document,"clarity","script","'+CTag+'");</script>\n')
#                     f.write(t2+'<link rel="stylesheet" href="'+rurl+'o.css')
#                     continue
#                 elif ls.startswith('<header'):
#                     f.write(l)
#                     if logo==True:
#                         f.write(t3+'<img src="'+bn.lower().replace(' ','-').replace('&','').replace("'",'')+'" alt="'+bn+' logo">\n')
#                     else:
#                         f.write(t3+'<h1>'+bn+'</h1>\n')
#                         if str(estYear) < 2000:
#                             f.write(t3+'<small>established in '+str(estYear)'</small>\n')
#                     if pn!='' or em!='' or sa!='':
#                         if pn!='':
#                             f.write(t3+'<img src="'+rurl+'phone-icon.svg" alt="show phone number icon" aria-label="show phone number" role="button" aria-pressed="false">\n')
#                         if em!='':
#                             if len(em) > 30:
#                                 f.write(t3+'<img src="'+rurl+'email-icon.svg" alt="show email address icon" aria-label="show email address" role="button" aria-pressed="false">\n')
#                             else:
#                                 f.write(t3+'<img src="'+rurl+'email-icon.svg" alt="show email address icon" aria-label="show email address" role="button" aria-pressed="false">\n')
#                         if sa!='':
#                             f.write(t3+'<img src="'+rurl+'directions-icon.svg" alt="show directions icon" aria-label="show phone number" role="button" aria-pressed="false">\n')
#                     continue
#                 elif ls.startswith('<main'):
#                     f.write(l)
#                     if logo==True:
#                         f.write(t3+'<h1>'+bn+'</h1>\n')
#                     f.write(t3+'<p id="p1">'+ps+ss+ars+'</p>\n')
#                     for i,pgS in enumerate(pgSs):
#                         f.write(t3+'<section>\n')
#                         f.write(t4+'<h2>'+pgS+'</h2>\n')
#                         f.write(t4+'<p></p>\n')
#                         f.write(t3+'</section>\n')
#                     if photos!=[]:
#                         f.write(m_imgs(dn))
#                     continue
#                 elif ls.startswith('<div id="m"'):
#                     f.write(l)
#                     if ad and sa!=[]:
#                         f.write(t3+'<div id="ma"><div id="maa" class="c"><address>'+street+'<br>'+city+','+state+','+zip+'</address><span>Get directions with Google Maps<img id="green-open-in-icon.svg"></span></div><div id="mas" class="c"></div>\n')
#                     elif sa!=[]:
#                         f.write(t3+'<div id="ma"><address>'+street+'<br>'+city+','+state+','+zip+'</address><a href="https://www.google.com/maps/dir/?api=1&destination='+street.replace(' ','+')+'%2C'+city+'%2C'+state+'+'+zip+'"><span>Get directions with Google Maps<img id="green-open-in-icon.svg"></span></a><a href=""></div>\n')
#                     else:
#                         lis=
#                         f.write(t3+'<div id="ma"><img src="map-showing-areas-served.png" alt="map showing areas served"><ul>'+lis+'</ul></div>\n')
#                     if pn!="":
#                         f.write(t3+'<a id="mp" href="tel:+1-"'+pn+'">'+pn+'</a>\n')
#                     if ea!="":
#                         f.write(t3+'<a id="me" href="mailto:" target="_blank"'+ea+'">'+ea+'</a>\n')
#                     if bh!=[]:
#                         f.write(t3+'<ul><li class="b">• Mon: '+Mon+'</li><li class="b">• Tue: '+Tue+'</li><li class="b">• Wed: '+Wed+'</li><li class="b">• Thu: '+Thu+'</li><li class="b">• Fri: '+Fri+'</li><li class="b">• Sat: '+Sat+'</li><li class="b">• Sun: '+Sun+'</li></ul>\n')
#                     if imgsStr!='':
#                         f.write(t3+'<img id="mi" alt="image enlarged after click">\n')
#                     continue
#                 else:
#                     f.write(l)
#     else:
#         for i,pg in enumerate(pgs):
#             a=[]
#             p=pg.lower()
#             p=p.replace(' ','-')
#             p=p.replace('&','and')
#             a.append(pg)
#             a.append(p)
#             pgs[i]=a
#         while("" in pgsS):
#             pgsS.remove("")
#         with open(lb+'m.html') as fi,open(dd+'index.html','w') as f:
#             for l in fi.readlines():
#                 ls=l.lstrip()
#                 if ls.startswith('<meta name="viewport"'):
#                     f.write(l)
#                     f.write(t2+'<title>'+bn+'</title>\n')
#                     f.write(t2+'<base href="'+base+'" target="_self">')
#                     f.write(t3+'<nav class="pt" aria-expanded="false">\n')
#                     f.write(footerStr)
#         for pg in pgs:
#             with open(lb+'m.html') as fi,open(dd+pg+'index.html','w') as f:
#                 for l in fi.readlines():
#                     ls=l.lstrip()
#                     if ls.startswith('<meta name="viewport"'):
#                         f.write(l)
#                         f.write(t2+'<title>'+bn+'</title>\n')
#                         f.write(t2+'<base href="'+base+'" target="_self">')
#                         f.write(t3+'<nav class="d">\n')
#                         f.write(t3+'<nav class="pt" aria-expanded="false">\n')
#                         for _pg in pgs:
#                             f.write(t4+'<a>Homepage</a>\n'
#                                 if _pg!=pg:
#                                     f.write(t4+'<a href="'+pg[1]+'">'+pg[0]+'</a>\n'
#                             f.write(t3+'</nav>\n')
#                         f.write(footerStr)
#     of(dn)
#     oc(dn)

def mf_footer(dn):
    dn=__import__(dn+'_d',fromlist=['i','l'])
    locals().update(i|l)
    while("" in ftr):
        ftr.remove("")
    if ftr!=[]:
        ftrInfo={'Yelp':'yelp-icon.png','Facebook':'facebook-icon-colored.png','Instagram':'instagram-icon.png','LinkedIn':'linkedin-icon.png','Twitter':'twitter-icon-colored-circle.svg','YouTube':'youtube-icon.','Better Business Bureau':'better-business-bureau-icon.png'}
    footerStr=t2+'<footer>\n'
    for footerSite in ftr:
        for k,v in ftrInfo.items():
            if k.lower() in footerSite:
                footerSiteHref=footerSite
                footerSiteName=k
                footerSiteIconSrc=v
                footerStr+=t3+'<a href="'+footerSiteHref+'" rel="noopener noreferrer" target="_blank"><img src="'+rurl+footerSiteIconSrc+'" alt="'+footerSiteName+' icon"></a></div>\n<small>a <a href="https://nothingfancywebsites.com" target="_blank">Nothing Fancy Website</a> by <span class="nw">Winners Webs</span></small>\n'
    footerStr+=t2+'</footer>\n'
    return footerStr

def mf_imgs(dn):
    startToSrc=t4+'<img class="w" src="'
    srcToAlt='" alt="'
    altToEnd='" loading="lazy">'
    imgsStr=''
    for f in os.listdir(dn):
        if f.endswith('.jpg' or '.png') and f[:-4]!='logo':
            imgsStr+=t4+startToSrc+f+srcToAlt+f[:-4].replace('-',' ')+altToEnd+'\n'
    return imgsStr

def miv(dn,vm):
    with open(p+'cData\\'+dn+'.txt','w') as f:
        document=Document(''+dn+'.docx')
        dP={}
        dS={}
        for i,row in enumerate(document.tables[0].rows):
            if i!=0:
                if row.cells[0].text!="":
                    dP.update({row.cells[0].text:{row.cells[1].text:row.cells[2].text}})
                if row.cells[3].text!="":
                    dS.update({row.cells[3].text:{row.cells[4].text:row.cells[5].text}})
        if dP!={}:
            f.write('products='+str(dP)+'\n')
        if dS!={}:
            f.write('services='+str(dS))
            for paragraph in document.paragraphs:
                l=paragraph.text
                if l.startswith('_') and l.count('"')==2:
                    if l[-1]=='"':
                        f.write(l+'\n')
                    else:
                        f.write(l[:l.rfind(' ')]+'\n')
                    continue
                elif l.startswith("Don't confirm"):
                    f.write(l.replace("Don't confirm",'')+'\n')
                    continue
                elif l.startswith('_') and l.count('"') > 2:
                    l=l.replace('"' ,"'" ,1)
                    l=l[:-1]+"'"
                    f.write(l+'\n')
    with open(p+'cData\\'+dn+'.txt') as fi,open(p+'cData\\'+dn+'-trimmed.txt','w') as f:
        for l in fi:
            if not l.endswith('""\n') or l.endswith('{}\n'):
                f.write(l)
    shutil.copyfile(p+'cData\\'+dn+'.txt',p+'py\\_cData\\'+dn+'.py')
    imp=__import__(dn,[])
    createCustomer()
    #cData || cDataObj?
    cData={'dn':dn,'dd':p+'dd\\'+dn+'\\','cData':p+'cData\\'+dn+'\\','lv':p+'lv\\'+dn+'\\','base':'https://'+dn+'/','vm':vm,'bn':imp._your_business_is_,'descr':imp._your_business_is_DESCRIBED_as_,'est':imp._the_DATE_you_were_ESTABLISHED_is_,'hist':imp._the_HISTORY_of_your_business_is_,'histMission':imp._the_HISTORY_or_VISION_of_your_business_is_,'mission':imp._the_MISSION_or_VISION_of_your_business_is_,'dnNew':imp._your_NEW_WEBSITE_will_be_named_,'dnCurrent':imp._your_CURRENT_WEBSITE_is_named_,'dnPast':imp._we_will_use_your_PAST_WEBSITE_name_,'Seo':imp._regarding_being_found_through_SEARCH_ENGINES_you_replied_,'ctTerm':imp._the_TERM_used_to_REFER_TO_your_CUSTOMERS_is_,'ctAvg':imp._your_AVERAGE_customerTerm_can_be_described_as_,'pdUsed':imp._PRODUCTS_proudly_USED_by_your_business_are_,'svUsed':imp._SERVICES_proudly_USED_by_your_business_are_,'refs':imp._you_sometimes_make_REFERRALS_to__,'ad':imp._your_business_ADDRESS_is_,'visFrom':imp._many_customerTerm_COME_FROM_,'visLocal':imp._regarding_most_customerTerm_being_LOCAL_you_replied_,'visWelcomed':imp._you_WELCOME_MORE_customerTerm_FROM_,'arsServed':imp._the_AREAS_you_SERVE_are_,'arsServedMost':imp._the_AREAS_you_SERVE_the_MOST_ARE_,'arsWelcomed':imp._you_WELCOME_MORE_customerTerm_IN_,'awardsCerts':imp._AWARDS_or_CERTIFICATIONS_relevant_to_your_business_are_,'ftr':{'Fb':imp._facebook_,'Ig':imp._instagram_,'Li':imp._linkedin_,'Tw':imp._twitter_,'Yt':imp._youtube_,'Yp':imp._yelp_,'Ot':imp._other_},'bh':[imp._sun_,imp._mon_,imp._tue_,imp._wed_,imp._thu_,imp._fri_,imp._sat_],'pn':imp._the_PHONE_NUMBER_used_for_customerTerm_is_,'ea':imp._the_EMAIL_ADDRESS_used_for_customerTerm_is_,'owners':imp._we_will_list_the_OWNERS_NAMES_as_,'pnA':imp._your_ADMINISTRATIVE_PHONE_NUMBER_is_,'eaA':imp._your_ADMINISTRATIVE_EMAIL_ADDRESS_is_,'pr':imp.products,'se':imp.services,'StId':StId}

    # if dnNew!="":
    #     dn=dnNew        
    # elif dnPast!="":
    #     dn=dnPast
    # else:
    #     dn=dnCurrent
    # pgs=[p2,p3,p4,p5]
    # while("" in pgs):
    #     pgs.remove("")
    with open(p+'py\\_cData\\'+dn+'_d.py','w') as f:
        f.write('i='+json.dumps(i)+'\nl='+json.dumps(l))

def nd(dn):
    dn=__import__(dn+'_d',fromlist=['i','l'])
    locals().update(i|l)
    for p,ds,fs in os.walk(lv):
        if fs.endswith('.html'):
            with open(fs+'index.html') as f:
                for i,l in enumerate(f.readlines()):
                    if i!=5:
                        f.write(l)
    with open(lv+'sitemap.xml') as f:
        f.write('<?xml version="1.0" encoding="UTF-8"?>\n<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n\t<url>\n\t\t<loc>'+base+'</loc>\n')
        if pgs!=[]:
            for pg in pgs:
                f.write('\t<url>\n\t\t<loc>'+base+pg+'</loc>\n\t</url>\n')
        f.write(t2+'\n\t</url>\n</urlset>')
    with open(lv+'robots.txt') as f:
        f.replace('Disallow','Allow')
        f.write('Sitemap: '+base+'/sitemap.xml')
    shutil.move(cData+dn+'-iv-trimmed.txt',p+'ar\\'+dn+'cData.txt')
    shutil.move(dd+'-iv-trimmed.txt',p+'ar\\'+dn+'cData-d.py')
    #mv d.py,rm dd

# def nd_cj(dn):
#     dn=__import__(dn+'_d',fromlist=['i','l'])
#     locals().update(i|l)
#     for p,ds,fs in os.walk(dn):
#         with open(p+'.html') as f:
#             allTagsAr=allClassesAr=allDataAr=allIdsAr=allScharsAr=[]
#             for l in f.readlines():
#                 tags=re.findall(r'\<([^>]*)\>',l)
#                     for tag in tags:
#                         #schars
#                         allTagsAr.append(tag)
#                 if 'class=' in l:
#                     classes=re.findall(r'class="([^"]*)"',l)
#                     for classAtt in classes:
#                         classAtt=classAtt.split(' ')
#                         for aClass in classAtt:
#                             allClassesAr.append(aClass)
#                 if 'id=' in l:
#                     ids=re.findall(r'id="([^"]*)"',l)
#                     for id in ids:
#                         allIdsAr.append(id)
#                 if ' data-' in l:
#                     dataKandV=re.findall(r' data-([^"]*"[^"]*)',l)
#                     dataKandV.split('="')
#                     i=[dataKAndV[0],dataKAndV[1]]
#                     allDataAr.append(i)
#             tagAr=[]
#             for i in allTagsAr:
#                 if i not in tagAr:
#                     tagAr.append(i)
#             # dataAr=[]
#             # for i in allDataAr:
#             #     if i not in dataAr:
#             #         dataAr.append(i)
#             classAr=[]
#             for i in allClassesAr:
#                 if i not in classAr:
#                     classAr.append(i)
#             idAr=[]
#             for i in allIdsAr:
#                 if i not in idAr:
#                     idAr.append(i)
#             scharsAr=[]
#             for i in allScharsAr:
#                 if i not in scharsAr:
#                     scharsAr.append(i)
#             return tagAr,classAr,dataAr,idAr
#         with open(dd+'m.css') as fi,open(dd+'style.css') as f:
#             for l in fi.readlines():
#                 ls=l.lstrip()
#                 # schar=['[',':','>',' ']
#                 if not ls.startswith('*/'):
#                     if ls.startswith('.'):
#                         className=ls.split('{')
#                         if className[0][1:] in classAr:
#                             f.write(l)
#                             continue
#                         else:
#                             print(l)
#                         continue
#                     elif ls.startswith('#'):
#                         idName=ls.split('{')
#                         if idName[0][1:] in idAr:
#                             f.write(l)
#                             continue
#                         else:
#                             print(l)
#                         continue
#                     elif ls.startswith('@media'):
#                         if not ls.endswith('{}\n'):
#                             f.write(l)
#                             continue
#                         else:
#                             print(l)
#                         continue
#                     elif ls.startswith('}'):
#                         f.write(l)
#                         continue
#                     else:
#                         selector=ls.split('{')
#                         selector=selector[0]
#                         for char in selector:
#                             if char in schar:
#                                 if char=='[':
#                                     dataKandV=re.findall(r'data-([^"]*"[^"]*)',selector)
#                                     dataKandV.split('="')
#                                     for i in dataAr:
#                                         if dataKAndV[0] in i[0]:
#                                             if dataKAndV[1] in i[1]:
#                                                 f.write(l)
#                                                 continue
#                                             else:
#                                                 print(l)
#                                             continue
#                                         else:
#                                             print(l)
#                                     # if dataKey in dataAr:
#                                     #     f.write(l)
#                                     #     continue
#                                     # else:
#                                     #     print(l)
#                                     # continue
#                                 elif char==':':

#                                     continue
#                                 elif char=='>':

#                                     continue
#                                 elif char==' ':
                                
#                                     continue
#                                 else:
#                                     print(l)
#                             else:
#                                 if selector in tagAr:
#                                     f.write(l)
#                                     continue
#                                 else:
#                                     print(l)
#                     #if pn!='' and sa!='' and em!='' and bh!=[]:
#                 else:
#                     print(l)
#         #shutil.remove(p+'m.css')
#         with open(dd+'m.js') as fi,open(dd+'function.js') as f:
#             for l in fi.readlines():
#                 ls=l.lstrip()
#                 if not ls.startswith('//') or not ls.startswith('/*'):
#                     if ls.startswith('const h='):
#                         if 'nav' in idAr:
#                             f.write(l)
#                             continue
#                         else:
#                             print(l)
#                         continue
#                     if ls.startswith('const m_='):
#                         if mi==True:
#                             f.write(l)
#                             continue
#                         else:
#                             print(l)
#                         continue
#                     if ls.startswith('const ci='):
#                         continue

#                         for i in dataAr:
#                             if i[0]=='ci':
#                                 f.write(l)
#                                 continue
#                             else:
#                                 print(l)
#                             continue

#                         for i in dataAr:
#                             if i[0]=='io':
#                                 f.write(l)
#                                 continue
#                             else:
#                                 print(l)
#                             continue

#                 else:
#                     print(l)

def oc(dn):
    for p,ds,fs in os.walk(dn):
        for f in fs:
            if f.endswith('.html'):
                webbrowser.get('chromium').open_new_tab(os.path.join(p,f))
                #change default .html app and open with code?

def of(dn):
    for p,ds,fs in os.walk(dn):
        for f in fs:
            if f.endswith('.html'):
                os.startfile(os.path.normpath(p+'\\'+f))

def imgs_html(dn):
    for p,ds,fs in os.walk(dn):
        p=os.path.join(p,'')
        imgsStr=''
        for f in fs:
            if f.endswith('.jpg'):
                if f[:12]!='map-showing-' and f[-8:-4]!='logo':
                    imgF=Image.open(p+f)
                    w,h=imgF.size
                    if w/h>=1.5:
                        cl='hil'
                    elif w/h>=1.2:
                        cl='hi'
                    elif h/w>=1.5:
                        cl='vil'
                    elif h/w>=1.2:
                        cl='vi'
                    else:
                        cl='si'
                    imgsStr+='\t\t\t\t<img class="'+cl+'" src="'+f+'" alt="'+f[:-4].replace('-',' ')+'" loading="lazy">'+'\n'
            elif f.endswith('.jpeg'):
                if f[:12]!='map-showing-' and f[-8:-5]!='logo':
                    imgF=Image.open(p+f)
                    w,h=imgF.size
                    if w/h>=1.5:
                        cl='hil'
                    elif w/h>=1.2:
                        cl='hi'
                    elif h/w>=1.5:
                        cl='vil'
                    elif h/w>=1.2:
                        cl='vi'
                    else:
                        cl='si'
                    imgsStr+='\t\t\t\t<img class="'+cl+'" src="'+f+'" alt="'+f[:-5].replace('-',' ')+'" loading="lazy">'+'\n'
        print(imgsStr)

#def in:

def imgs_resize(dn):
    for p,ds,fs in os.walk(dn):
        p=os.path.join(p,'')
        if p.endswith('ir\\'):
            for f in fs:
                with Image.open(os.path.join(p+f)) as image:
                    w,h=image.size
                    if w>1440 or h>1440:
                        if (w>h):
                            nw=1440
                            nh=int(h/w*1440)
                            image=image.resize((nw,nh),Image.ANTIALIAS)
                        else:
                            nw=int(w/h*1440)
                            nh=1440
                            image=image.resize((nw,nh),Image.ANTIALIAS)
                    image=ImageOps.exif_transpose(image)
                    image.save(p[:-3]+f,quality=40,optimize=True)
        # for p,ds,fs in os.walk(dn):
        # p=os.path.join(p,'')
        # if not os.path.exists(p+'d'):
        #     os.makedirs(p+'d')
        # if not p.endswith('d\\'):
        #     for f in fs:
        #         f=f.lower()
        #         if f.endswith('.jpg') or f.endswith('.jpeg'):
        #             with Image.open(os.path.join(p+f)) as image:
        #                 w,h=image.size
        #                 if w>1440 or h>1440:
        #                     moveFlag=True
        #                 else:
        #                     moveFlag=False
        #             if moveFlag==True:
        #                 shutil.move(p+f,p+'d\\'+f)

def rm(dn):
    shutil.move(cData+dn+'-trimmed.txt.',p+'ar\\dnp\\'+dn+'-trimmed.txt')
    #shutil.move(cData+dn+'.txt',p+'ar\\dnp\\'+dn+'.txt')
    shutil.remove(cData+dn+'.txt')
    shutil.remove(p+'py\\_cData\\'+dn+'.py')

#def rma():
    #html
    #css

# def ph(dn,**GTag=False,**CTag=False):
#     dn=__import__(dn+'_d',fromlist=['i','l'])
#     locals().update(i|l)
#     #update ph with logo
#     if isFile(lv+'index.html'):
#         with open(lv+'index.html') as fi,open(lv+'w.html','w') as fo:
#             for l in fi.readlines():
#                 ls=l.lstrip()
#                 if ls.startswith('<title'):
#                     fo.write(l)
#                     fo.write(t2+'<link rel="icon" type="image/png" href="'+base+'favicon.png">\n')
#                     continue
#                 elif l.startswith(t3+'span {'):
#                     fo.write(l)
#                     fo.write(t3+'img {max-width:30%; margin:auto auto 40px;}\n')
#                     fo.write(t3+'@media screen and (max-width:768px) {\n'+t4+'img {max-width:50%;}\n'+t3+'}\n')
#                     continue
#                 elif l.startswith(t4+'span {'):
#                     fo.write(l)
#                     fo.write(t4+'img {max-width:90%;}')
#                     continue
#                 elif ls.startswith('<body'):
#                     fo.write(l)
#                     fo.write(t2+'<img src="' base+'logo'+logoExt+'" alt="'+bn+' logo>\n'
#                     continue
#                 else:
#                     fo.write(l)
#         shutil.copyfile(lv+'index.html',p+'dd\\ph-preLogo-'+dn+'.html')
#         os.remove(lv+'index.html')
#         os.rename(lv+'w.html',lv+'index.html')
#         up(dn)
#     #mkph w/ logo
#     elif logo is True:#detect logo file
#         with open(lb+'ph.html') as fi,open(lv+'index.html','w') as f:
#             for l in fi.readlines():
#                 ls=l.lstrip()
#                 if ls.startswith('<title'):
#                     l=l.replace('~bn',bn)
#                     f.write(l)
#                     continue
#                 elif l.startswith(t3+'span {'):
#                     fo.write(l)
#                     fo.write(t3+'img {max-width:30%; margin:auto auto 40px;}\n')
#                     fo.write(t3+'@media screen and (max-width: 768px) {\n'+t4+'img {max-width:50%;}\n'+t3+'}\n')
#                     continue
#                 elif l.startswith(t4+'span {'):
#                     fo.write(l)
#                     fo.write(t4+'img {max-width:90%;}')
#                     #more css
#                     if GTag!=False:
#                         f.write(t2+'<script async src="https://www.googletagmanager.com/gtag/js?id=G-'+GTag+'"></script><script>window.dataLayer=window.dataLayer || []; function gtag(){dataLayer.push(arguments);} gtag(\'js\',new Date()); gtag(\'config\',\'G-'+GTag+'\');</script>\n')
#                     if CTag!=False:
#                         f.write(t2+'<script>(function(c,l,a,r,i,t,y){c[a]=c[a]||function(){(c[a].q=c[a].q||[]).push(arguments)};t=l.createElement(r);t.async=1;t.src="https://www.clarity.ms/tag/"+i;y=l.getElementsByTagName(r)[0];y.parentNode.insertBefore(t,y);})(window,document,"clarity","script","'+CTag+'");</script>\n')
#                         f.write(l)
#                         continue
#                 elif ls.startswith('<em'):
#                     if currentWebsiteName=="" and pastWebsiteName!=availableWebsiteName:
#                         l=l.replace('the all new<br>','')
#                     l=l.replace('~dn',dn)
#                     f.write(l)
#                     continue
#                 else:
#                     f.write(l)
#         shutil.copyfile(lb+'robots.txt',lv+'robots.txt')
#         up(dn)
#     #mkph w/o logo
#     else:
#         md()
#         with open(lb+'ph.html') as fi,open(lv+'index.html','w') as f:
#             for l in fi.readlines():
#                 ls=l.lstrip()
#                 if ls.startswith('<title'):
#                     l=l.replace('~bn',bn)
#                     f.write(l)
#                     continue
#                 #if ls.startswith('<script'):
#                     #l=l.replace('~GTag',GTag,2)
#                     #l=l.replace('~CTag',CTag)
#                     #f.write(l)
#                 elif ls.startswith('<em'):
#                     if currentWebsiteName=="" and pastWebsiteName!=availableWebsiteName:
#                         l=l.replace('the all new<br>','')
#                     l=l.replace('~dn',dn)
#                     f.write(l)
#                     continue
#                 else:
#                     f.write(l)
#         shutil.copyfile(lb+'robots.txt',lv+'robots.txt')
#         up(dn)

# def up(dn):
#     dn=__import__(dn+'_d',fromlist=['i','l'])
#     locals().update(i|l)
#     kpr=''
#     client=paramiko.SSHClient()
#     client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
#     client.connect(ip,username='',pkey=kpr)
#     if dnCurrentStaysDn==False:
#         if '/var/www/'+dn not in client.sftp.listdir('/var/www/'+dn):
#             stdin,stdout,stderr=client.exec_command('bash add-site '+dn)
#             for line in stdout:
#                 print('... '+line.strip('\n'))
#     client.close()
#     with pysftp.Connection(ip,username='',private_key=kpr) as sftp:
#         sftp.put_r(lv+dn,'/var/www/'+dn,preserve_mtime=True)
#         sftp.close()





# #unminify html
# ts+='\t'
# [:-2]

# #unminify css
# replace('}','}\n')
# if ('@'):
#     mediaFlag=True
#     replace('{','{\n')
#     '\t'
#     replace('}','}\n')

#unm js
#one character before each line written from m.js + \n
